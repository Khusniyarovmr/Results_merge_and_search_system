import os
import re
from datetime import datetime
import sqlite3
from CONSTANTS import DATA_BASE

from PyQt5 import QtCore
import docx


class Searching(QtCore.QThread):
    is_inn = 0
    interested_text = ''
    data_for_table = QtCore.pyqtSignal(list)
    spisokList = QtCore.pyqtSignal(list)
    finished = QtCore.pyqtSignal()
    status_mes = QtCore.pyqtSignal(str)

    def __init__(self, *args):
        super().__init__()
        self.working = True

    def run(self):
        try:
            self.data_for_table.emit(self.search_text(self.interested_text))
        except TypeError:
            pass
        finally:
            self.finished.emit()

    def podgotovka_znacheniy(self, text_for_change):
        text_from_input = text_for_change.lower()
        search_string = []
        text_from_input = text_from_input.replace('  ', ' ')
        text_from_input = re.sub("^\s+|\n|\r|\s+$", '', text_from_input)
        startIndex = text_from_input.find('\"')
        kol_vo_probelov = text_from_input.count(' ')
        if text_from_input.isnumeric():  # Все введенные символы - числовые. По русски: ввели ИНН
            search_string.append(text_from_input)
            self.is_inn = 1
        elif startIndex != -1 and kol_vo_probelov != 2:  # Нашли в введенных параметрах символ кавычек. Значит ищем по наименованию компании.
            search_string.append(text_from_input)
            endIndex = text_from_input.find('\"', startIndex + 1)
            if startIndex != -1 and endIndex != -1:  # i.e. both quotes were found
                search_string.append(text_from_input[startIndex + 1:endIndex])
        elif kol_vo_probelov == 2 and startIndex == -1:  # Если в веденных параметрах два пробела и нет кавычек. то ФИО.
            search_string.append(text_from_input)
            fiotext = re.split('\s+', text_from_input)
            search_string.append(
                fiotext[0] + 'ым ' + fiotext[1] + 'ом ' + fiotext[2] + 'ем')  # Ивановым Иваном Васильевичем
            search_string.append(fiotext[0] + 'у ' + fiotext[1] + 'у ' + fiotext[2] + 'у')  # Иванову Ивану Васильевичу
            search_string.append(fiotext[0] + 'а ' + fiotext[1] + 'а ' + fiotext[2] + 'а')  # Иванова Ивана Васильевича
            search_string.append(fiotext[0][:-1] + 'ой ' + fiotext[1][:-1] + 'ой ' + fiotext[2][
                                                                                     :-1] + 'ой')  # Ивановой Татьяной Васильевной
            search_string.append(fiotext[0] + ' ' + fiotext[1][:1] + '.' + fiotext[2][:1] + '.')  # Иванов И.В.
            search_string.append(fiotext[0] + 'а' + ' ' + fiotext[1][:1] + '.' + fiotext[2][:1] + '.')  # Иванова И.В.
            search_string.append(fiotext[0] + 'ым' + ' ' + fiotext[1][:1] + '.' + fiotext[2][:1] + '.')  # Ивановым И.В.
            search_string.append(fiotext[0] + 'у' + ' ' + fiotext[1][:1] + '.' + fiotext[2][:1] + '.')  # Иванову И.В.
            search_string.append(
                fiotext[0][:-1] + 'ой' + ' ' + fiotext[1][:1] + '.' + fiotext[2][:1] + '.')  # Ивановой Т.В.
        else:  # Если введенные параметры не числовые и в них нет кавычек, значит это или ФИО или наименование.
            search_string.append(text_from_input)
        self.spisokList.emit(search_string)
        return search_string

    def search_text(self, text_for_search):
        while self.working:
            if text_for_search == '': return
            t0 = datetime.now()
            self.status_mes.emit(str('Поиск начался в ' + str(datetime.now())))
            search_string = self.podgotovka_znacheniy(text_for_search)
            t1 = datetime.now() - t0
            t_now = datetime.now()
            self.status_mes.emit(str('(1/9) Завершилась функция подготовки значений ' + str(t1.seconds) + ' секунд'))
            if not self.poisk_cancel(self.stoped()): return
            # Запускаем поиск.
            # Поиск в Заключениях и в отчете АБ. Используем WORD.
            search_result_in_zakl = self.poisk_v_zakl(search_string)
            t2 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(str('(2/9) Завершилась функция поиска в заключениях ' + str(
                t2.seconds) + ' с. Запуск поиска в отчетах АБ'))
            if not self.poisk_cancel(self.stoped()): return
            
            search_result_in_otchet_ab = self.poisk_v_otchetah_ab(search_string)
            t2 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(str('(3/9) Завершилась функция поиска в заключениях ' + str(
                t2.seconds) + ' с. Запуск поиска в целевом'))
            if not self.poisk_cancel(self.stoped()): return
            
            # Поиск в Целевом, Реестре попавших в выборку. Используем Excel.
            # search_result_in_celevoe = self.poisk_v_celevom(search_string)
            search_result_in_celevoe = self.poisk_v_xlsx(search_string, 'file_2')
            t3 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(
                str('(4/9) Завершилась функция поиска в целевом ' + str(t3.seconds) + ' с. Запуск поиска в выборке'))
            if not self.poisk_cancel(self.stoped()): return
            
            # search_result_in_viborka = self.poisk_v_viborke(search_string)
            search_result_in_viborka = self.poisk_v_xlsx(search_string, 'file_4')
            t4 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(
                str('(5/9) Завершилась функция поиска в выборке ' + str(t4.seconds) + ' с. Запуск поиска в клиентах'))
            if not self.poisk_cancel(self.stoped()): return
            
            # search_result_in_master_file = self.poisk_v_viborke(search_string)
            search_result_in_master_file = self.poisk_v_xlsx(search_string, 'file_8')
            t4 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(
                str('(6/9) Завершилась функция поиска в выборке ' + str(t4.seconds) + ' с. Запуск поиска в клиентах'))
            if not self.poisk_cancel(self.stoped()): return

            # Поиск в клиентах банка, реестре сделок, реестре бенифициаров. Используем Postgresql
            search_result_in_clients = self.poisk_v_bank_clients(search_string)
            t5 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(
                str('(7/9) Завершилась функция поиска в клиентах ' + str(t5.seconds) + ' с. Запуск поиска в счетах'))
            if not self.poisk_cancel(self.stoped()): return
            
            client_string = []
            for i in range(0, len(search_result_in_clients)):
                client_string.append([search_result_in_clients[i][j] for j in range(2)])
            search_result_in_client_accounts = self.poisk_v_clients_account(client_string)
            t6 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(str('(8/9) Завершилась функция поиска в счетах ' + str(
                t6.seconds) + ' с. Запуск поиска в оценке заемщиков'))
            if not self.poisk_cancel(self.stoped()): return
            
            search_result_in_ocenka_zaemchikov = self.poisk_v_ocenke_zaemchikov(search_result_in_clients)
            t7 = datetime.now() - t_now
            t_now = datetime.now()
            self.status_mes.emit(
                str('(9/9) Завершилась функция поиска в оценке заемщиков ' + str(t7.seconds) + ' с.'))
            if not self.poisk_cancel(self.stoped()): return
            
            if search_result_in_client_accounts:
                client_zaemchik = {str(id_client_acc[0]) + '^' + str(id_client_acc[1]): 'Клиент' for id_client_acc in
                                   search_result_in_client_accounts}
                for i in range(0, len(search_result_in_client_accounts)):
                    if search_result_in_client_accounts[i][2][:2] == '45':
                        client_zaemchik[str(search_result_in_client_accounts[i][0]) + '^' + str(
                            search_result_in_client_accounts[i][1])] = 'Клиент/Заемщик'

                ocenka_zaemchika = {str(id_inn_oc[0]) + '^' + str(id_inn_oc[2]): '' for id_inn_oc in
                                    search_result_in_ocenka_zaemchikov}
                for i in range(0, len(search_result_in_ocenka_zaemchikov)):
                    ocenka_zaemchika[
                        str(search_result_in_ocenka_zaemchikov[i][0]) + '^' + str(
                            search_result_in_ocenka_zaemchikov[i][2])] = \
                        search_result_in_ocenka_zaemchikov[i][3]
            if not self.poisk_cancel(self.stoped()): return
            
            name_vseh_bankov = self.reestr_bankov()
            dic_bank_names = {id_name_bank[0]: id_name_bank[1] for id_name_bank in name_vseh_bankov}
            string_for_table = []
            spisok_id_bankov = []
            if not self.poisk_cancel(self.stoped()): return
            for zak in search_result_in_zakl: spisok_id_bankov.append(zak[0])
            for otch in search_result_in_otchet_ab: spisok_id_bankov.append(otch[0])
            for celev in search_result_in_celevoe: spisok_id_bankov.append(celev[0])
            for vib in search_result_in_viborka: spisok_id_bankov.append(vib[0])
            for cl in search_result_in_clients: spisok_id_bankov.append(cl[0])
            for mf in search_result_in_master_file: spisok_id_bankov.append(mf[0])
            spisok_id_bankov = list(set(spisok_id_bankov))
            if not self.poisk_cancel(self.stoped()): return
            
            proverka = {}
            for i in range(len(spisok_id_bankov)):
                for j in range(len(search_result_in_clients)):
                    a = ['', '', '', '', '', '', '', '', '', '', '', '', '']
                    if spisok_id_bankov[i] == search_result_in_clients[j][0]:
                        a[0] = dic_bank_names[search_result_in_clients[j][0]]
                        a[1] = self.ZAO_OOO_OAO(search_result_in_clients[j][2])
                        a[2] = search_result_in_clients[j][3]
                        a[3] = ' '.join(re.sub('\s+', '', search_result_in_clients[j][k]) for k in range(4, 6))
                        a[4] = client_zaemchik[
                            str(search_result_in_clients[j][0]) + '^' + str(search_result_in_clients[j][1])]
                        a[5] = ocenka_zaemchika.get(
                            str(search_result_in_clients[j][0]) + '^' + str(search_result_in_clients[j][3]))
                        a[12] = search_result_in_clients[j][1]
                        string_for_table.append(a)
                        proverka[dic_bank_names[spisok_id_bankov[i]]] = ''

            for i in range(len(spisok_id_bankov)):
                if proverka.get(dic_bank_names[spisok_id_bankov[i]]) == None:
                    a = [dic_bank_names[spisok_id_bankov[i]], '', '', '', '', '', '', '', '', '', '', '', '', '']
                    string_for_table.append(a)
                    proverka[dic_bank_names[spisok_id_bankov[i]]] = ''
            if not self.poisk_cancel(self.stoped()): return

            for k in range(len(string_for_table)):
                for j in range(len(search_result_in_viborka)):
                    if dic_bank_names[search_result_in_viborka[j][0]] == string_for_table[k][0]:
                        string_for_table[k][9] = search_result_in_viborka[j][1]
            if not self.poisk_cancel(self.stoped()): return

            for k in range(len(string_for_table)):
                for j in range(len(search_result_in_celevoe)):
                    if dic_bank_names[search_result_in_celevoe[j][0]] == string_for_table[k][0]:
                        string_for_table[k][10] = search_result_in_celevoe[j][1]
            if not self.poisk_cancel(self.stoped()): return

            for k in range(len(string_for_table)):
                for j in range(len(search_result_in_master_file)):
                    if dic_bank_names[search_result_in_master_file[j][0]] == string_for_table[k][0]:
                        string_for_table[k][11] = search_result_in_master_file[j][1]
            if not self.poisk_cancel(self.stoped()): return

            for k in range(len(string_for_table)):
                for j in range(len(search_result_in_zakl)):
                    if dic_bank_names[search_result_in_zakl[j][0]] == string_for_table[k][0]:
                        string_for_table[k][7] = search_result_in_zakl[j][1]
            if not self.poisk_cancel(self.stoped()): return

            for k in range(len(string_for_table)):
                for j in range(len(search_result_in_otchet_ab)):
                    if dic_bank_names[search_result_in_otchet_ab[j][0]] == string_for_table[k][0]:
                        string_for_table[k][8] = search_result_in_otchet_ab[j][1]
            if not self.poisk_cancel(self.stoped()): return

            self.working = False
            t8 = datetime.now() - t0
            self.status_mes.emit(str('Процесс поиска завершен. Прошло ' + str(t8.seconds) + ' секунд'))
        return string_for_table

    def stoped(self):
        return self.working

    def poisk_cancel(self, fact):
        if not fact:
            self.status_mes.emit(str('Поиск отменён!'))
            self.finished.emit()
            return None
        return 'Zero'

    def poisk_v_zakl(self, spisok):
        con_to_DB = sqlite3.connect(DATA_BASE)
        cur = con_to_DB.cursor()
        cur.execute("SELECT bank_id, file_1 FROM bank_information WHERE bank_information.file_1 != ''")
        file_zakluchenia_string = cur.fetchall()
        cur.close()
        con_to_DB.close()
        itogi_poiska_zak = []
        for i in range(len(file_zakluchenia_string)):
            if file_zakluchenia_string[i][1] != '':
                file_name = os.path.basename(file_zakluchenia_string[i][1])
                full_file_path = os.path.join('C:/NEW/Файлы_заключений_для_системы/Заключения/Поиск', file_name)
                result_find_in_file = []
                if not self.stoped():
                    self.status_mes.emit(str('Поиск отменён!'))
                    self.finished.emit()
                    return
                doc = docx.Document(full_file_path)
                for paragraph in doc.paragraphs:
                    for s_name in spisok:
                        if paragraph.text.lower().find(s_name) != -1:
                            result_find_in_file.append(file_zakluchenia_string[i])
                itogi_poiska_zak.extend(result_find_in_file)
        itogi_poiska_zak = list(set(itogi_poiska_zak))
        return itogi_poiska_zak

    def poisk_v_otchetah_ab(self, spisok):
        con_to_DB = sqlite3.connect(DATA_BASE)
        cur = con_to_DB.cursor()
        cur.execute("SELECT bank_id, file_7 FROM bank_information WHERE bank_information.file_7 != ''")
        file_zakluchenia_string = cur.fetchall()
        cur.close()
        con_to_DB.close()
        itogi_poiska_otchetah_ab = []
        for i in range(len(file_zakluchenia_string)):
            if file_zakluchenia_string[i][1] != '':
                file_name = os.path.basename(file_zakluchenia_string[i][1])
                full_file_path = os.path.join('C:/NEW/Файлы_заключений_для_системы/Отчет_АБ/Поиск', file_name)
                result_find_in_file = []
                if not self.stoped():
                    self.status_mes.emit(str('Поиск отменён!'))
                    self.finished.emit()
                    return
                doc = docx.Document(full_file_path)
                for paragraph in doc.paragraphs:
                    for s_name in spisok:
                        if paragraph.text.lower().find(s_name) != -1:
                            result_find_in_file.append(file_zakluchenia_string[i])
                itogi_poiska_otchetah_ab.extend(result_find_in_file)
        itogi_poiska_otchetah_ab = list(set(itogi_poiska_otchetah_ab))
        return itogi_poiska_otchetah_ab

    def poisk_v_xlsx(self, spisok, file):
        con_to_DB = sqlite3.connect(DATA_BASE)
        cur = con_to_DB.cursor()
        cur.execute(
            """SELECT bank_id, """ + file + """ FROM bank_information WHERE bank_information.""" + file + """ != ''""")
        files_string = cur.fetchall()
        cur.close()
        con_to_DB.close()
        itogi_poiska_xlsx = []
        for i in range(len(files_string)):
            if files_string[i][1] != '':
                dirname, file_name = os.path.split(files_string[i][1])
                dirname = str(dirname) + '/Поиск'
                full_file_path = os.path.join(dirname, file_name.replace('xlsx', 'csv'))
                if not self.stoped():
                    self.status_mes.emit(str('Поиск отменён!'))
                    self.finished.emit()
                    return
                cur_file = open(full_file_path, 'r', encoding='UTF-8')
                for row in cur_file:
                    for s_name in spisok:
                        if row.lower().find(s_name) != -1:
                            itogi_poiska_xlsx.append(files_string[i])
        itogi_poiska_xlsx = list(set(itogi_poiska_xlsx))
        return itogi_poiska_xlsx

    def poisk_v_bank_clients(self, spisok):
        con_to_DB = sqlite3.connect(DATA_BASE)
        cur = con_to_DB.cursor()
        cur.execute("SELECT bank_id, id_client, name, inn, doc_seria, doc_number FROM clients")
        bank_clients_string = cur.fetchall()
        cur.close()
        con_to_DB.close()
        itogi_poiska_clients = []
        if self.is_inn == 1:  # Смотрим ИНН
            for i in range(0, len(bank_clients_string)):
                if not self.stoped():
                    self.status_mes.emit(str('Поиск отменён!'))
                    self.finished.emit()
                    return
                for s_name in spisok:
                    if s_name.lower() in bank_clients_string[i][3].lower():
                        itogi_poiska_clients.append(bank_clients_string[i])
            self.is_inn = 0
        else:  # Смотрим наименование
            for i in range(0, len(bank_clients_string)):
                if not self.stoped():
                    self.status_mes.emit(str('Поиск отменён!'))
                    self.finished.emit()
                    return
                for s_name in spisok:
                    if s_name.lower() in bank_clients_string[i][2].lower():
                        itogi_poiska_clients.append(bank_clients_string[i])
        itogi_poiska_clients = list(set(itogi_poiska_clients))
        # print('\n'.join(str(itogi_poiska_clients[i]) for i in range(0, len(itogi_poiska_clients))))
        return itogi_poiska_clients

    def poisk_v_clients_account(self, spisok):
        # print(spisok)
        if spisok == []: return
        con_to_DB = sqlite3.connect(DATA_BASE)
        cur = con_to_DB.cursor()
        itogi_poiska_client_accounts = []
        for k in range(len(spisok)):
            if not self.stoped():
                self.status_mes.emit(str('Поиск отменён!'))
                self.finished.emit()
                return
            cur.execute("SELECT bank_id, id_client, account, acc_name FROM accounts WHERE accounts.bank_id = '" + str(
                spisok[k][0]) + "' and accounts.id_client = '" + str(spisok[k][1]) + "'")
            clients_account_string = cur.fetchall()
            if clients_account_string:
                # Смотрим совпадение по bank_id и id_client
                for i in range(0, len(clients_account_string)):
                    itogi_poiska_client_accounts.append(clients_account_string[i])
                # print('\n'.join(str(itogi_poiska_client_accounts[i]) for i in range(0, len(itogi_poiska_client_accounts))))
        cur.close()
        con_to_DB.close()
        return itogi_poiska_client_accounts

    def poisk_v_ocenke_zaemchikov(self, spisok):
        con_to_DB = sqlite3.connect(DATA_BASE)
        cur = con_to_DB.cursor()
        cur.execute("SELECT * FROM reestr_popavshih_v_viborku")
        ocenka_zaemchikov_string = cur.fetchall()
        cur.close()
        con_to_DB.close()
        itogi_poiska_v_ocenke_zaemchikov = []
        # Смотрим совпадение по bank_id и id_client
        for i in range(0, len(ocenka_zaemchikov_string)):
            for j in range(0, len(spisok)):
                if not self.stoped():
                    self.status_mes.emit(str('Поиск отменён!'))
                    self.finished.emit()
                    return
                if ocenka_zaemchikov_string[i][0] == spisok[j][0] and ocenka_zaemchikov_string[i][2] == spisok[j][3]:
                    itogi_poiska_v_ocenke_zaemchikov.append(ocenka_zaemchikov_string[i])
                elif ocenka_zaemchikov_string[i][0] == spisok[j][0] and ocenka_zaemchikov_string[i][2] == ' '.join(
                        re.sub('\s+', '', spisok[j][k]) for k in range(4, 6)):
                    itogi_poiska_v_ocenke_zaemchikov.append(ocenka_zaemchikov_string[i])

        # print('\n'.join(str(itogi_poiska_v_ocenke_zaemchikov[i]) for i in range(0, len(itogi_poiska_v_ocenke_zaemchikov))))
        return itogi_poiska_v_ocenke_zaemchikov

    def reestr_bankov(self):
        con_to_DB = sqlite3.connect(DATA_BASE)
        cur = con_to_DB.cursor()
        cur.execute("SELECT id, small_name FROM banks")
        reestr_bankov = cur.fetchall()
        cur.close()
        con_to_DB.close()
        return reestr_bankov

    def ZAO_OOO_OAO(self, text):
        itog = re.sub(r"(\b[ЗзАаКк][а-яА-ЯёЁ]{5,}\s+\b[АаКкЦц][а-яА-ЯёЁ]*\s+\b[ОоБб][а-яА-ЯёЁ]*)", 'ЗАО', text,
                      re.IGNORECASE)  # Закрытое акционерное общество
        if 'ЗАО' not in itog:
            itog = re.sub(r"(\b[ОоБбЩщ][а-яА-ЯёЁ]{5,}\s+[Сс]\s+\b[ОоГгРр][а-яА-ЯёЁ]*\s+\b[ОоТтВв][а-яА-ЯёЁ]*)", 'ООО',
                          text,
                          re.IGNORECASE)  # Общество с ограниченной ответственностью
            if 'ООО' not in itog:
                itog = re.sub(r"(\b[ОоТтКк][а-яА-ЯёЁ]{5,}\s+\b[АаКкЦц][а-яА-ЯёЁ]*\s+\b[ОоБб][а-яА-ЯёЁ]*)", 'ОАО', text,
                              re.IGNORECASE)  # Открытое акционерное общество
                if 'ОАО' not in itog:
                    itog = re.sub(r"(\b[ПпУуБб][а-яА-ЯёЁ]{6,}\s+\b[АаКкЦц][а-яА-ЯёЁ]*\s+\b[ОоБб][а-яА-ЯёЁ]*)", 'ПАО',
                                  text,
                                  re.IGNORECASE)  # Публичное акционерное общество
                    if 'ПАО' not in itog:
                        itog = re.sub(r"(\b[АаКкЦц][а-яА-ЯёЁ]*\s+\b[ОоБб][а-яА-ЯёЁ]*)", 'АО', text,
                                      re.IGNORECASE)  # Акционерое общество
        return itog
